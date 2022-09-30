# -*- coding: utf-8 -*-
# @Time    : 2022/8/7 15:48
# @Author  : 银尘
# @FileName: pdf_util.py
# @Software: PyCharm
# @Email   ：liwudi@liwudi.fun

import os
import json
import logging
from time import perf_counter
from multiprocessing import Pool, cpu_count
import fitz
from docx import Document
from pdf2docx.page.Page import Page
from pdf2docx.page.Pages import Pages
import io
from PaperCrawlerUtil.crawler_util import verify_rule
from common_util import *
import pdfplumber
from PyPDF2 import PdfFileWriter, PdfFileReader
from pdf2docx import Converter


# check PyMuPDF>=1.19.x
if list(map(int, fitz.VersionBind.split("."))) < [1, 19, 0]:
    raise SystemExit("PyMuPDF>=1.19.0 is required for pdf2docx.")


class ConversionException(Exception):
    pass


class MakedocxException(ConversionException):
    pass


class MyPdf2DocxConverter(Converter):
    '''The ``PDF`` to ``docx`` converter.

        * Read PDF file with ``PyMuPDF`` to get raw layout data page by page, including text,
          image, drawing and its properties, e.g. boundary box, font, size, image width, height.
        * Analyze layout in document level, e.g. page header, footer and margin.
        * Parse page layout to docx structure, e.g. paragraph and its properties like indentaton,
          spacing, text alignment; table and its properties like border, shading, merging.
        * Finally, generate docx with ``python-docx``.
        '''

    def __init__(self, pdf_file: str, password: str = None):
        '''Initialize fitz object with given pdf file path.

        Args:
            pdf_file (str): pdf file path.
            password (str): Password for encrypted pdf. Default to None if not encrypted.
        '''
        # fitz object
        self.filename_pdf = pdf_file
        self.password = str(password or '')
        self._fitz_doc = fitz.Document(pdf_file)

        # initialize empty pages container
        self._pages = Pages()

    @property
    def fitz_doc(self):
        return self._fitz_doc

    @property
    def pages(self):
        return self._pages

    def close(self):
        self._fitz_doc.close()

    @property
    def default_settings(self):
        '''Default parsing parameters.'''
        return {
            'debug': False,  # plot layout if True
            'ocr': 0,  # ocr status: 0 - no ocr; 1 - to do ocr; 2 - ocr-ed pdf
            'ignore_page_error': True,  # not break the conversion process due to failure of a certain page if True
            'multi_processing': False,  # convert pages with multi-processing if True
            'cpu_count': 0,  # working cpu count when convert pages with multi-processing
            'min_section_height': 20.0,  # The minimum height of a valid section.
            'connected_border_tolerance': 0.5,  # two borders are intersected if the gap lower than this value
            'max_border_width': 6.0,  # max border width
            'min_border_clearance': 2.0,  # the minimum allowable clearance of two borders
            'float_image_ignorable_gap': 5.0,  # float image if the intersection exceeds this value
            'page_margin_factor_top': 0.5,  # [0,1] reduce top margin by factor
            'page_margin_factor_bottom': 0.5,  # [0,1] reduce bottom margin by factor
            'shape_min_dimension': 2.0,  # ignore shape if both width and height is lower than this value
            'max_line_spacing_ratio': 1.5,  # maximum line spacing ratio: line spacing / line height
            'line_overlap_threshold': 0.9,  # [0,1] delete line if the intersection to other lines exceeds this value
            'line_break_width_ratio': 0.5,
            # break line if the ratio of line width to entire layout bbox is lower than this value
            'line_break_free_space_ratio': 0.1,
            # break line if the ratio of free space to entire line exceeds this value
            'line_separate_threshold': 5.0,  # two separate lines if the x-distance exceeds this value
            'new_paragraph_free_space_ratio': 0.85,
            # new paragraph if the ratio of free space to line height exceeds this value
            'lines_left_aligned_threshold': 1.0,  # left aligned if d_x0 of two lines is lower than this value (Pt)
            'lines_right_aligned_threshold': 1.0,  # right aligned if d_x1 of two lines is lower than this value (Pt)
            'lines_center_aligned_threshold': 2.0,
            # center aligned if delta center of two lines is lower than this value
            'clip_image_res_ratio': 4.0,  # resolution ratio (to 72dpi) when cliping page image
            'min_svg_gap_dx': 15.0,  # merge adjacent vector graphics if the horizontal gap is less than this value
            'min_svg_gap_dy': 2.0,  # merge adjacent vector graphics if the vertical gap is less than this value
            'min_svg_w': 2.0,  # ignore vector graphics if the bbox width is less than this value
            'min_svg_h': 2.0,  # ignore vector graphics if the bbox height is less than this value
            'extract_stream_table': False,  # don't consider stream table when extracting tables
            'parse_lattice_table': True,  # whether parse lattice table or not; may destroy the layout if set False
            'parse_stream_table': True,  # whether parse stream table or not; may destroy the layout if set False
            'delete_end_line_hyphen': False  # delete hyphen at the end of a line
        }

    # -----------------------------------------------------------------------
    # Parsing process: load -> analyze document -> parse pages -> make docx
    # -----------------------------------------------------------------------

    def parse(self, start: int = 0, end: int = None, pages: list = None, **kwargs):
        '''Parse pages in three steps:
        * open PDF file with ``PyMuPDF``
        * analyze whole document, e.g. page section, header/footer and margin
        * parse specified pages, e.g. paragraph, image and table

        Args:
            start (int, optional): First page to process. Defaults to 0, the first page.
            end (int, optional): Last page to process. Defaults to None, the last page.
            pages (list, optional): Range of page indexes to parse. Defaults to None.
            kwargs (dict, optional): Configuration parameters.
        '''
        return self.load_pages(start, end, pages) \
            .parse_document(**kwargs) \
            .parse_pages(**kwargs)

    def load_pages(self, start: int = 0, end: int = None, pages: list = None):
        '''Step 1 of converting process: open PDF file with ``PyMuPDF``,
        especially for password encrypted file.

        Args:
            start (int, optional): First page to process. Defaults to 0, the first page.
            end (int, optional): Last page to process. Defaults to None, the last page.
            pages (list, optional): Range of page indexes to parse. Defaults to None.
        '''
        log(self._color_output('[1/4] Opening document...'))

        # encrypted pdf ?
        if self._fitz_doc.needs_pass:
            if not self.password:
                raise ConversionException(f'Require password for {self.filename_pdf}.')

            elif not self._fitz_doc.authenticate(self.password):
                raise ConversionException('Incorrect password.')

        # initialize empty pages
        num = len(self._fitz_doc)
        self._pages.reset([Page(id=i, skip_parsing=True) for i in range(num)])
        log("there are {} pages need to be converted".format(str(num)))
        # set pages to parse
        page_indexes = self._page_indexes(start, end, pages, num)
        for i in page_indexes:
            self._pages[i].skip_parsing = False

        return self

    def parse_document(self, **kwargs):
        '''Step 2 of converting process: analyze whole document, e.g. page section,
        header/footer and margin.'''
        log(self._color_output('[2/4] Analyzing document...'))

        self._pages.parse(self.fitz_doc, **kwargs)
        return self

    def parse_pages(self, **kwargs):
        '''Step 3 of converting process: parse pages, e.g. paragraph, image and table.'''
        log(self._color_output('[3/4] Parsing pages...'))

        pages = [page for page in self._pages if not page.skip_parsing]
        num_pages = len(pages)
        p_bar = process_bar(unit="page", final_prompt="finish")
        p_bar.process(0, 1, num_pages)
        for i, page in enumerate(pages, start=1):
            p_bar.process(i, 1, num_pages)
            pid = page.id + 1
            log('({}/{}) Page {}'.format(str(i), str(num_pages), str(pid)))
            try:
                page.parse(**kwargs)
            except Exception as e:
                if not kwargs['debug'] and kwargs['ignore_page_error']:
                    log('Ignore page {} due to parsing page error: {}'.format(str(pid), str(e)))
                else:
                    raise ConversionException(f'Error when parsing page {pid}: {e}')

        return self

    def make_docx(self, docx_filename=None, **kwargs):
        '''Step 4 of converting process: create docx file with converted pages.

        Args:
            docx_filename (str): docx filename to write to.
            kwargs (dict, optional): Configuration parameters.
        '''
        log(self._color_output('[4/4] Creating pages...'))

        # check parsed pages
        parsed_pages = list(filter(
            lambda page: page.finalized, self._pages
        ))
        if not parsed_pages:
            raise ConversionException('No parsed pages. Please parse page first.')

        # docx file to convert to
        filename = docx_filename or f'{self.filename_pdf[0:-len(".pdf")]}.docx'
        if os.path.exists(filename): os.remove(filename)

        # create page by page
        docx_file = Document()
        num_pages = len(parsed_pages)
        for i, page in enumerate(parsed_pages, start=1):
            if not page.finalized: continue  # ignore unparsed pages
            pid = page.id + 1
            log('({}/{}) Page {}'.format(str(i), str(num_pages), str(pid)))
            try:
                page.make_docx(docx_file)
            except Exception as e:
                if not kwargs['debug'] and kwargs['ignore_page_error']:
                    log('Ignore page {} due to making page error: {}'.format(str(pid), str(e)))
                else:
                    raise MakedocxException(f'Error when make page {pid}: {e}')

        # save docx
        docx_file.save(filename)

    # -----------------------------------------------------------------------
    # Store / restore parsed results
    # -----------------------------------------------------------------------

    def store(self):
        '''Store parsed pages in dict format.'''
        return {
            'filename': os.path.basename(self.filename_pdf),
            'page_cnt': len(self._pages),  # count of all pages
            'pages': [page.store() for page in self._pages if page.finalized],  # parsed pages only
        }

    def restore(self, data: dict):
        '''Restore pages from parsed results.'''
        # init empty pages if necessary
        if not self._pages:
            num = data.get('page_cnt', 100)
            self._pages.reset([Page(id=i, skip_parsing=True) for i in range(num)])

        # restore pages
        for raw_page in data.get('pages', []):
            idx = raw_page.get('id', -1)
            self._pages[idx].restore(raw_page)

    def serialize(self, filename: str):
        '''Write parsed pages to specified JSON file.'''
        with open(filename, 'w', encoding='utf-8') as f:
            f.write(json.dumps(self.store(), indent=4))

    def deserialize(self, filename: str):
        '''Load parsed pages from specified JSON file.'''
        with open(filename, 'r') as f:
            data = json.load(f)
        self.restore(data)

    # -----------------------------------------------------------------------
    # high level methods, e.g. convert, extract table
    # -----------------------------------------------------------------------

    def debug_page(self, i: int, docx_filename: str = None, debug_pdf: str = None, layout_file: str = None, **kwargs):
        '''Parse, create and plot single page for debug purpose.

        Args:
            i (int): Page index to convert.
            docx_filename (str): docx filename to write to.
            debug_pdf (str): New pdf file storing layout information. Default to add prefix ``debug_``.
            layout_file (str): New json file storing parsed layout data. Default to ``layout.json``.
        '''
        # include debug information
        # fitz object in debug mode: plot page layout
        # file path for this debug pdf: demo.pdf -> debug_demo.pdf
        path, filename = os.path.split(self.filename_pdf)
        if not debug_pdf: debug_pdf = os.path.join(path, f'debug_{filename}')
        if not layout_file: layout_file = os.path.join(path, 'layout.json')
        kwargs.update({
            'debug': True,
            'debug_doc': fitz.Document(),
            'debug_filename': debug_pdf
        })

        # parse and create docx
        self.convert(docx_filename, pages=[i], **kwargs)

        # layout information for debugging
        self.serialize(layout_file)

    def convert(self, docx_filename: str = None, start: int = 0, end: int = None, pages: list = None, **kwargs):
        """Convert specified PDF pages to docx file.

        Args:
            docx_filename (str, optional): docx filename to write to. Defaults to None.
            start (int, optional): First page to process. Defaults to 0, the first page.
            end (int, optional): Last page to process. Defaults to None, the last page.
            pages (list, optional): Range of page indexes. Defaults to None.
            kwargs (dict, optional): Configuration parameters. Defaults to None.

        Refer to :py:meth:`~pdf2docx.converter.Converter.default_settings` for detail of
        configuration parameters.

        .. note::
            Change extension from ``pdf`` to ``docx`` if ``docx_file`` is None.

        .. note::
            * ``start`` and ``end`` is counted from zero if ``--zero_based_index=True`` (by default).
            * Start from the first page if ``start`` is omitted.
            * End with the last page if ``end`` is omitted.

        .. note::
            ``pages`` has a higher priority than ``start`` and ``end``. ``start`` and ``end`` works only
            if ``pages`` is omitted.

        .. note::
            Multi-processing works only for continuous pages specified by ``start`` and ``end`` only.
        """
        t0 = perf_counter()
        log('Start to convert {}'.format(self.filename_pdf))
        settings = self.default_settings
        settings.update(kwargs)

        # input check
        if pages and settings['multi_processing']:
            raise ConversionException('Multi-processing works for continuous pages '
                                      'specified by "start" and "end" only.')

        # convert page by page
        if settings['multi_processing']:
            self._convert_with_multi_processing(docx_filename, start, end, **settings)
        else:
            self.parse(start, end, pages, **settings).make_docx(docx_filename, **settings)

        log('Terminated in {}.'.format(str(perf_counter() - t0)))

    def extract_tables(self, start: int = 0, end: int = None, pages: list = None, **kwargs):
        '''Extract table contents from specified PDF pages.

        Args:
            start (int, optional): First page to process. Defaults to 0, the first page.
            end (int, optional): Last page to process. Defaults to None, the last page.
            pages (list, optional): Range of page indexes. Defaults to None.
            kwargs (dict, optional): Configuration parameters. Defaults to None.

        Returns:
            list: A list of parsed table content.
        '''
        # parsing pages first
        settings = self.default_settings
        settings.update(kwargs)
        self.parse(start, end, pages, **settings)

        # get parsed tables
        tables = []
        for page in self._pages:
            if page.finalized: tables.extend(page.extract_tables(**settings))

        return tables

    def _convert_with_multi_processing(self, docx_filename: str, start: int, end: int, **kwargs):
        '''Parse and create pages based on page indexes with multi-processing.

        Reference:

            https://pymupdf.readthedocs.io/en/latest/faq.html#multiprocessing
        '''
        # make vectors of arguments for the processes
        cpu = min(kwargs['cpu_count'], cpu_count()) if kwargs['cpu_count'] else cpu_count()
        prefix = 'pages'  # json file writing parsed pages per process
        vectors = [(i, cpu, start, end, self.filename_pdf, self.password,
                    kwargs, f'{prefix}-{i}.json') for i in range(cpu)]

        # start parsing processes
        pool = Pool()
        pool.map(self._parse_pages_per_cpu, vectors, 1)

        # restore parsed page data
        for i in range(cpu):
            filename = f'{prefix}-{i}.json'
            if not os.path.exists(filename): continue
            self.deserialize(filename)
            os.remove(filename)

        # create docx file
        self.make_docx(docx_filename, **kwargs)

    @staticmethod
    def _parse_pages_per_cpu(vector):
        '''Render a page range of a document.

        Args:
            vector (list): A list containing required parameters.
                * 0  : segment number for current process
                * 1  : count of CPUs
                * 2,3: whole pages range to process
                * 4  : pdf filename
                * 5  : password for encrypted pdf
                * 6  : configuration parameters
                * 7  : json filename storing parsed results
        '''
        # recreate the arguments
        idx, cpu, s, e, pdf_filename, password, kwargs, json_filename = vector

        # open pdf to get page count: all pages are marked to parse temporarily
        # since don't know which pages to parse for this moment
        cv = Converter(pdf_filename, password)
        cv.load_pages()

        # the specified pages to process
        e = e or len(cv.fitz_doc)
        all_indexes = range(s, e)
        num_pages = len(all_indexes)

        # page segment processed by this cpu
        m = int(num_pages / cpu)
        n = num_pages % cpu
        seg_size = m + int(idx < n)
        seg_from = (m + 1) * idx + min(n - idx, 0)
        seg_to = min(seg_from + seg_size, num_pages)
        page_indexes = [all_indexes[i] for i in range(seg_from, seg_to)]

        # now, mark the right pages
        for page in cv.pages: page.skip_parsing = True
        for i in page_indexes:
            cv.pages[i].skip_parsing = False

        # parse pages and serialize data for further processing
        cv.parse_document(**kwargs) \
            .parse_pages(**kwargs) \
            .serialize(json_filename)
        cv.close()

    @staticmethod
    def _page_indexes(start, end, pages, pdf_len):
        '''Parsing arguments.'''
        if pages:
            indexes = [int(x) for x in pages]
        else:
            end = end or pdf_len
            s = slice(int(start), int(end))
            indexes = range(pdf_len)[s]

        return indexes

    @staticmethod
    def _color_output(msg):
        return f'\033[1;36m{msg}\033[0m'


def pdf2docx(pdf_path: str, word_path: str, end_pages: int = None,
             start_pages: str = None, need_log: bool = True) -> None:
    """
    转换pdf 到word文件，可以自动识别是文件夹还是单个文件，其中word_path表示的生成的word的文件夹，不论是
    单个还是文件夹批量转换，这个值都是文件夹
    :param need_log: 是否需要日志
    :param pdf_path: pdf的路径
    :param word_path: 用于存放word的路径，必须是文件夹路径
    :param end_pages: 结束页码
    :param start_pages: 开始页码
    :return:
    """
    file_list = []
    file = True
    count = 0
    if os.path.isfile(pdf_path):
        file_list.append(pdf_path)
        if need_log:
            log("转换文件{}开始".format(pdf_path))
    else:
        file_list.extend(getAllFiles(pdf_path))
        if need_log:
            log("获取文件夹{}文件成功".format(pdf_path))
        file = False
    for ele in file_list:
        if ele.endswith(".pdf"):
            try:
                cv = MyPdf2DocxConverter(ele)
                if start_pages is None:
                    start_pages = 0
                if end_pages is None:
                    end_pages = len(cv.pages)
                cv.convert(local_path_generate(word_path, suffix=".docx"), start=start_pages, end=end_pages)
                count = count + 1
                log("总计pdf文件个数{}，已经完成{}".format(len(file_list), count))
            except Exception as e:
                log(string="转换失败文件{},{}".format(ele, e), print_file=sys.stderr)
            finally:
                cv.close()


def get_para_from_one_pdf(path: str, begin_tag: list = None,
                          end_tag: list = None, ranges: tuple = (0, 1)) -> str:
    """
        用来从pdf文件中获取一些文字，可以通过设置开始或者结束标志，以及页码范围获取自己想要的内容
        如果是文件夹，则直接遍历文件夹中所有的PDF，返回所有符合的字符串，同时可以设置分隔符
        :param path: pdf path
        :param begin_tag: the tag which will begin from this position to abstract text
        :param end_tag: the tag which will end util this position to abstract text
        :param ranges: which pages you want to abstract
        :return: the string
    """
    txt = ""
    try:
        with pdfplumber.open(path) as pdf:
            if len(pdf.pages) >= 0:
                left_range = ranges[0]
                right_range = ranges[1] + 1
                if right_range >= len(pdf.pages):
                    right_range = len(pdf.pages)
                if left_range >= len(pdf.pages):
                    left_range = 0
                for i in range(left_range, right_range):
                    txt = txt + pdf.pages[i].extract_text()
                if len(begin_tag) == 0 and len(end_tag) == 0:
                    txt = txt
                elif len(begin_tag) == 0 and len(end_tag) > 0:
                    ele = ""
                    for e in end_tag:
                        if txt.find(e) >= 0:
                            ele = e
                            break
                    if len(ele) > 0:
                        txt = txt.split(ele)[0]
                elif len(begin_tag) > 0 and len(end_tag) == 0:
                    ele = ""
                    for e in begin_tag:
                        if txt.find(e) >= 0:
                            ele = e
                            break
                    if len(ele) > 0:
                        txt = txt.split(ele)[1]
                elif len(begin_tag) > 0 and len(end_tag) > 0:
                    ele1 = ""
                    ele2 = ""
                    for e1 in begin_tag:
                        if txt.find(e1) >= 0:
                            ele1 = e1
                            break
                    for e2 in end_tag:
                        if txt.find(e2) >= 0:
                            ele2 = e2
                            break
                    if len(ele1) > 0 and len(ele2) > 0:
                        txt = txt.split(ele1)[1].split(ele2)[0]
        pdf.close()
        return txt
    except Exception as e:
        log(string="打开PDF异常：{}".format(e), print_file=sys.stderr)
        return txt


def get_para_from_pdf(path: str, begin_tag: list = None, end_tag: list = None, ranges: tuple = (0, 1),
                      split_style: str = "===", valid_threshold: int = 0, need_log: bool = True) -> str:
    """
    用来从pdf文件中获取一些文字，可以通过设置开始或者结束标志，以及页码范围获取自己想要的内容
    如果是文件夹，则直接遍历文件夹中所有的PDF，返回所有符合的字符串，同时可以设置分隔符
    :param need_log: 是否需要日志
    :param valid_threshold: decide whether paragraph digest success
    :param path: pdf path
    :param begin_tag: the tag which will begin from this position to abstract text
    :param end_tag: the tag which will end util this position to abstract text
    :param ranges: which pages you want to abstract
    :param split_style: split style which will used to split string of each file of directory
    :return: the string
    """
    if begin_tag is None:
        begin_tag = []
    if end_tag is None:
        end_tag = ["1. Introduction", "1. introduction",
                   "Introduction", "introduction",
                   "1.摘要", "1. 摘要", "1.", "1"]
    txt = ""
    valid_count = 0
    sum_count = 0
    file_list = []
    if os.path.isfile(path):
        file_list.append(path)
    else:
        file_list.extend(getAllFiles(path))
    for ele in file_list:
        if ele.endswith("pdf"):
            tem = get_para_from_one_pdf(ele, begin_tag, end_tag, ranges)
            txt = txt + tem + "\n"
            txt = txt + get_split(style=split_style) + get_split(style="\n", lens=3)
            if len(tem) > valid_threshold:
                valid_count = valid_count + 1
                if need_log:
                    log("有效抽取文件：{}".format(ele))
            else:
                log(string="抽取文件疑似失败：{}".format(ele), print_file=sys.stderr)
            sum_count = sum_count + 1
        else:
            sum_count = sum_count + 1
            log(string="错误：{}不是PDF文件".format(ele), print_file=sys.stderr)
    log("总计抽取了文件数量：{}，其中有效抽取（>{}）数量：{}".format(sum_count, valid_threshold, valid_count))
    return txt


def getAllFiles(target_dir: str) -> list:
    """
    遍历文件夹
    :param target_dir: 遍历的文件夹
    :return: 所有文件的名称
    """
    files = []
    if len(target_dir) == 0:
        log(string="文件路径为空", print_file=sys.stderr)
        return files
    try:
        listFiles = os.listdir(target_dir)
    except Exception as e:
        log(string="打开文件夹{}异常：{}".format(target_dir, e), print_file=sys.stderr)
        return files
    for i in range(0, len(listFiles)):
        path = os.path.join(target_dir, listFiles[i])
        if os.path.isdir(path):
            files.extend(getAllFiles(path))
        elif os.path.isfile(path):
            files.append(path)
    return files


class sub_func_write_pdf(threading.Thread):

    def __init__(self, out_path: str, out_stream: io.BufferedWriter, out_pdf: PdfFileWriter) -> None:
        threading.Thread.__init__(self)
        self.res = False
        self.out_path = out_path
        self.out_stream = out_stream
        self.output = out_pdf

    def run(self) -> None:
        super().run()
        try:
            self.output.write(self.out_stream)
            self.res = True
            self.out_stream.close()
        except Exception as e:
            log(string=threading.currentThread().getName() + "写PDF出现异常：{}".format(e), print_file=sys.stderr)
            self.res = False

    def getRes(self):
        return self.res

    def raiseException(self):
        raise ThreadStopException()


def getSomePagesFromOnePDF(path: str, out_path: str, page_range: tuple or list,
                           need_log: bool = True, timeout: float = 20) -> bool:
    """
        从给定的文件路径中，截取指定的页面，保存到给定输出目录中
        :param path: 文件夹或者文件路径
        :param page_range: 截取的范围，如果是元组，则连续截取[a,b]的页面，注意从0开始，如果是列表，
        则按照列表给出的信息截取
        :param out_path:指定输出的目录
        :param need_log: 是否需要打印日志
        :param timeout: 线程结束时间，防止转换过程中，线程无意义等待
        :return: 返回布尔值，True表示成功转换
        """

    if len(path) == 0:
        log(string="路径参数为空，返回错误", print_file=sys.stderr)
        return False
    if type(page_range) != tuple and type(page_range) != list:
        log(string="页码范围有误，返回错误", print_file=sys.stderr)
        return False
    output = None
    pdf_file = None
    pdf_pages_len = None
    try:
        output = PdfFileWriter()
        pdf_file = PdfFileReader(open(path, "rb"))
        pdf_pages_len = pdf_file.getNumPages()
    except Exception as e:
        log(string="打开文件异常：{}".format(e), print_file=sys.stderr)
        return False
    iters = None
    if type(page_range) == tuple:
        new_page_range = ()
        if len(page_range) == 0:
            log(string="页码范围不明确，返回错误", print_file=sys.stderr)
            return False
        elif len(page_range) == 1:
            log("使用范围截取，但只有一个参数，结束参数默认为最大值")
            new_page_range = (page_range[0], pdf_pages_len - 1)
        elif len(page_range) > 2:
            log("使用范围参数，但参数数量过多，截取两个")
            new_page_range = (page_range[0], page_range[1])
        else:
            new_page_range = (page_range[0], page_range[1])
        # check tuple legal or not
        a = []
        for k in new_page_range:
            if k < 0:
                a.append(0)
            elif k > pdf_pages_len:
                a.append(pdf_pages_len)
            a.append(k)
        new_page_range = (a[0], a[1])
        if new_page_range[1] < new_page_range[0]:
            new_page_range = (new_page_range[1], new_page_range[0])
        iters = range(new_page_range[0], new_page_range[1])
    else:
        # 去重
        page_range = list(set(page_range))
        for k in page_range:
            '''@todo: 完善verify_rule()'''
            if not (0 <= k <= pdf_pages_len - 1):
                log(string="范围参数有错", print_file=sys.stderr)
                return False
        iters = page_range
    for i in iters:
        output.addPage(pdf_file.getPage(i))
    try:
        outputStream = open(out_path, "wb")
        sub = sub_func_write_pdf(out_path, outputStream, output)
        sub.setDaemon(True)
        sub.start()
        sub.join(timeout=timeout)
        success = sub.getRes()
        try:
            sub.raiseException()
        except Exception as e:
            if need_log:
                log(string=e, print_file=sys.stderr)
        if need_log or not success:
            log(("从文件{}截取页面到{}成功".format(path, out_path)) if success else ("从文件{}截取页面到{}失败".format(path, out_path)))
        return True
    except Exception as e:
        log(string="写文件出错：{}".format(e), print_file=sys.stderr)
        return False
    finally:
        pdf_file.stream.close()


def getSomePagesFromFileOrDirectory(path: str, page_range: tuple or list, out_directory: str = "",
                                    need_log: bool = True, timeout: float = 20, need_bar: bool = True) -> None:
    """
    从给定的文件夹或者文件路径中，截取指定的页面，保存到给定输出目录中
    :param need_bar: 是否需要使用进度条代替日志显示，此项为True时，覆盖need_log参数
    :param path: 文件夹或者文件路径
    :param page_range: 截取的范围，如果是元组，则连续截取[a,b]的页面，注意从0开始，如果是列表，
    则按照列表给出的信息截取
    :param out_directory:指定输出的目录
    :param need_log: 是否需要打印日志
    :param timeout: 线程结束时间，防止转换过程中，线程无意义等待
    :return:
    """

    count = 0
    sum = 0
    need_log = need_log if not need_bar else False
    if os.path.isfile(path):
        p = process_bar(total=1, desc="文件截取进度：", final_prompt="文件截取完成")
        p.process(0, 1, 1)
        sum = 1
        if getSomePagesFromOnePDF(path, local_path_generate(out_directory, need_log=need_log), page_range, need_log):
            count = count + 1
            p.process(1, 1, 1)
    else:
        files = getAllFiles(path)
        total = 0
        for f in files:
            if f.endswith(".pdf"):
                total = total + 1
        p = process_bar(total=total, desc="文件截取进度：", final_prompt="文件截取完成")
        p.process(0, 1, total)
        for k in files:
            if not k.endswith(".pdf"):
                log(string="文件{}不是PDF文件".format(k), print_file=sys.stderr)
                continue
            sum = sum + 1
            if getSomePagesFromOnePDF(k, local_path_generate(out_directory, need_log=need_log), page_range, need_log,
                                      timeout):
                count = count + 1
                p.process(count, 1, total)
    if need_log:
        log("总计待截取文件：{}，成功：{}".format(str(sum), str(count)))


def cooperatePdfWithLimit(files: list, page_range: tuple or list = None, out_path: str = "",
                          need_log: bool = True, timeout: float = -1, group_id: str = "",
                          need_group: bool = True) -> bool:
    """
    合并列表的PDF文件到指定目录
    :param files: 文件列表
    :param page_range: 合并的页码范围，如果是元组，则连续截取[a,b]的页面，注意从0开始，如果是列表，
    则按照列表给出的信息截取
    :param out_path: 输出文件的目录，如果不给定则默认自动生成
    :param need_log: 是否需要日志
    :param timeout: 线程等待时间，如果该值为-1则一直等到线程执行完毕，否则等待时间间隔之后，将结束线程，
    未合并完成的文件则自动失败
    :param group_id: 分组号，会自动添加在给定的文件名上，如果给定目标文件名为空，则自动生成文件名
    :param need_group: 是否需要分组
    :return:
    """
    output = None
    if need_group:
        out_path = list(out_path)
        temp = []
        for p in range(len(out_path) - 4):
            temp.append(out_path[p])
        out_path = "".join(temp)
        out_path = (out_path + "-groupid-" +group_id + ".pdf") if len(out_path) != 0 else local_path_generate("")
    else:
        out_path = out_path if len(out_path) != 0 else local_path_generate("")
    count = 0
    try:
        output = PdfFileWriter()
    except Exception as e:
        log(string="打开PDF写文件工具失败：{}".format(e), print_file=sys.stderr)
    file_readers = []
    for file in files:
        reader = None
        if not file.endswith(".pdf"):
            log(string="文件{}不是PDF文件，略过".format(file), print_file=sys.stderr)
            continue
        try:
            if file == out_path:
                continue
            reader = PdfFileReader(open(file, "rb"))
        except Exception as e:
            log(string="打开文件{}失败{}".format(file, e), print_file=sys.stderr)
            continue
        pdf_pages_len = reader.getNumPages()
        iters = None
        if type(page_range) == tuple:
            new_page_range = ()
            if len(page_range) == 0:
                if need_log:
                    log("默认全部合并，因为范围为空")
                new_page_range = (0, pdf_pages_len)
            elif len(page_range) == 1:
                if need_log:
                    log("使用范围截取，但只有一个参数，结束参数默认为最大值")
                new_page_range = (page_range[0], pdf_pages_len)
            elif len(page_range) > 2:
                if need_log:
                    log("使用范围参数，但参数数量过多，截取两个")
                new_page_range = (page_range[0], page_range[1])
            else:
                new_page_range = (page_range[0], page_range[1])
            iters = range(new_page_range[0], new_page_range[1])
        else:
            page_range = list(set(page_range))
            for k in page_range:
                try:
                    if not (verify_rule({0: GREATER_AND_EQUAL, pdf_pages_len - 1: LESS_THAN_AND_EQUAL}, float(k))):
                        log(string="范围参数有错", print_file=sys.stderr)
                        return False
                except Exception as e:
                    log(string="参数范围输入格式错误：{}".format(e), print_file=sys.stderr)
                    return False
            iters = page_range
        for i in iters:
            output.addPage(reader.getPage(i))
        file_readers.append(reader)
        count = count + 1
    try:
        outputStream = open(out_path, "wb")
    except Exception as e:
        log(string="打开文件异常：{}".format(e), print_file=sys.stderr)
        return False
    sub = sub_func_write_pdf(out_path, outputStream, output)
    # sub.setDaemon(True)
    sub.start()
    if timeout <= 0:
        sub.join()
    else:
        sub.join(timeout=timeout)
    flag = sub.getRes()
    try:
        sub.raiseException()
    except ThreadStopException as e:
        if need_log:
            log("结束线程")
    for read in file_readers:
        read.stream.close()
    if flag:
        if need_log:
            log("合并文件到{}成功，共计{}文件，合并总数{}".format(out_path, str(len(files)), str(count)))
        return True
    else:
        log(string="合并失败", print_file=sys.stderr)
        return False


def cooperatePdf(path: str, page_range: tuple or list = None, out_path: str = "",
                 need_log: bool = True, timeout: float = -1, group_number: int = 50,
                 need_group: bool = True, need_processbar: bool = True) -> None:
    """
    合并文件夹中所有的PDF文件到指定目录
    :param need_processbar: 是否需要使用进度条代替日志显示，此项为True时，覆盖need_log参数
    :param path: 文件夹路径
    :param page_range: 合并的页码范围，如果是元组，则连续截取[a,b]的页面，注意从0开始，如果是列表，
    则按照列表给出的信息截取，如果此项值不给定，则默认全部合并
    :param out_path: 输出文件的目录，如果不给定则默认自动生成
    :param need_log: 是否需要日志
    :param timeout: 线程等待时间，如果该值为-1则一直等到线程执行完毕，否则等待时间间隔之后，将结束线程，
    未合并完成的文件则自动失败
    :param group_number: 分组时每组的合并数量
    :param need_group: 是否需要分组合并
    :return:
    """
    if len(path) == 0:
        log("给定路径为空，合并结束：{}".format(path))
    elif os.path.isfile(path):
        log("给定的是文件路径，合并结束：{}".format(path))
    if page_range is None:
        page_range = ()
    files = getAllFiles(path)
    if need_group:
        if need_processbar:
            p = process_bar(desc="文件合并进度：", total=len(files), final_prompt="文件合并完成")
            p.process(0, group_number, len(files))
        i: int = 0
        group_id: int = 0
        file_group = []
        while i < len(files):
            if (i != 0 and ((i % group_number) == 0)) or (i == len(files) - 1):
                if i == len(files) - 1:
                    file_group.append(files[i])
                cooperatePdfWithLimit(file_group, page_range, out_path, need_log, timeout, str(group_id))
                group_id = group_id + 1
                file_group.clear()
                if need_processbar:
                    p.process(0, group_number, len(files))
            file_group.append(files[i])
            i = i + 1
    else:
        if need_processbar:
            p = process_bar(desc="文件合并进度：", total=len(files), final_prompt="文件合并完成")
            p.process(0, group_number, len(files))
        cooperatePdfWithLimit(files, page_range, out_path, need_log, timeout, need_group=False)
        if need_processbar:
            p.process(0, len(files), len(files))